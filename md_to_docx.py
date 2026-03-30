import re
from docx import Document
from docx.shared import Pt

def convert_md_to_docx(md_file, docx_file):
    doc = Document()
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Headers
        if line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        
        # Bold text like **Text**
        elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or line.startswith('4. '):
             p = doc.add_paragraph(line, style='List Number')
        
        # Regular text
        else:
            # Handle inline bold **text**
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == "__main__":
    convert_md_to_docx('D:\\moonbit\\MoonAgent_Proposal.md', 'D:\\moonbit\\MoonAgent_Proposal.docx')
