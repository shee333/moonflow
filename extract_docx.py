from docx import Document

doc_path = "MoonFlow基于MoonBit的Dify式AI智能体编排平台_韦旭.docx"
doc = Document(doc_path)

print("=" * 80)
print("MOONFLOW项目文档内容")
print("=" * 80)

for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)
        print()

print("\n" + "=" * 80)
print("文档中的表格数量:", len(doc.tables))
print("=" * 80)
